import asyncio


from pypdf import PdfReader
from docx import Document


from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession
from sqlalchemy.orm import sessionmaker

from .worker import celery_app
from .config import get_settings
from . import models
from .s3_client import download_file_from_s3


settings = get_settings()


async def process_file_async(file_id: int):

    engine = create_async_engine(settings.sqlalchemy_database_url)
    async_session_factory = sessionmaker(
        engine, class_=AsyncSession, expire_on_commit=False
    )

    async with async_session_factory() as session:

        file = await session.get(models.File, file_id)
        if not file:
            print(f"Файл с ID {file_id} не найден в базе данных.")
            return

        print(f"Найден файл: {file.filename}, s3_path: {file.s3_path}")

        file_stream = download_file_from_s3(file.s3_path)
        if not file_stream:
            print(f"Не удалось скачать файл {file.s3_path} из S3.")
            return

        metadata_to_update = {}

        if file.filename.lower().endswith(".pdf"):
            try:
                reader = PdfReader(file_stream)
                meta = reader.metadata
                metadata_to_update["page_count"] = len(reader.pages)
                metadata_to_update["title"] = meta.title
                metadata_to_update["author"] = meta.author
                metadata_to_update["creator_tool"] = meta.creator
            except Exception as e:
                print(f"Ошибка при парсинге PDF {file.filename}: {e}")

        elif file.filename.lower().endswith(".docx"):
            try:
                document = Document(file_stream)
                meta = document.core_properties
                metadata_to_update["author"] = meta.author
                metadata_to_update["title"] = meta.title
                metadata_to_update["created_date"] = str(meta.created)
                metadata_to_update["paragraph_count"] = len(document.paragraphs)
                metadata_to_update["table_count"] = len(document.tables)
            except Exception as e:
                print(f"Ошибка при парсинге DOCX {file.filename}: {e}")

        if metadata_to_update:
            print(f"Обновляем метаданные для файла {file_id}: {metadata_to_update}")
            for key, value in metadata_to_update.items():
                setattr(file, key, value)
            await session.commit()
            print(f"Метаданные для файла {file_id} успешно сохранены.")
        else:
            print(f"Для файла {file_id} не найдено метаданных для извлечения.")

    await engine.dispose()


@celery_app.task
def extract_metadata(file_id: int):
    print(f"Запущена задача extract_metadata для файла ID: {file_id}")

    try:
        asyncio.run(process_file_async(file_id))
        result_message = f"Обработка файла {file_id} завершена успешно."
    except Exception as e:
        print(f"Произошла ошибка при обработке файла {file_id}: {e}")
        result_message = f"Ошибка при обработке файла {file_id}."

    return result_message
